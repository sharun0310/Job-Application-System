import fitz  # PyMuPDF
import docx
import easyocr
import io
import logging
import numpy as np
from PIL import Image

logger = logging.getLogger(__name__)


class TextExtractor:
    def __init__(self):
        try:
            self.reader = easyocr.Reader(["en"], gpu=False)
        except Exception as e:
            logger.error(f"Failed to initialize EasyOCR: {e}")
            self.reader = None

    def extract_from_pdf(self, file_bytes: bytes) -> str:
        text = ""
        try:
            doc = fitz.open(stream=file_bytes, filetype="pdf")
            for page in doc:
                page_text = page.get_text()
                if page_text and page_text.strip():
                    text += page_text + "\n"
                else:
                    # OCR fallback for scanned PDF pages
                    pix = page.get_pixmap()
                    img_bytes = pix.tobytes("png")
                    text += self.extract_from_image(img_bytes) + "\n"
            return text.strip()
        except Exception as e:
            logger.error(f"PDF Extraction error: {e}")
            return ""

    def extract_from_docx(self, file_bytes: bytes) -> str:
        text = ""
        try:
            doc = docx.Document(io.BytesIO(file_bytes))
            for para in doc.paragraphs:
                if para.text:
                    text += para.text + "\n"
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text:
                            text += cell.text + "\n"
            return text.strip()
        except Exception as e:
            logger.error(f"DOCX Extraction error: {e}")
            return ""

    def extract_from_image(self, file_bytes: bytes) -> str:
        if not self.reader:
            return ""
        try:
            image = Image.open(io.BytesIO(file_bytes)).convert("RGB")
            img_np = np.array(image)
            result = self.reader.readtext(img_np, detail=0)
            return " ".join(result).strip()
        except Exception as e:
            logger.error(f"Image OCR error: {e}")
            return ""


text_extractor = TextExtractor()
